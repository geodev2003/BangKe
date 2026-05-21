"""
Điền dữ liệu vào file Word mẫu template.docx rồi dùng LibreOffice xuất PDF.
File mẫu: /app/templates/template.docx (hoặc bất kỳ .docx nào trong /app/templates/)
"""
import copy, os, shutil, subprocess, tempfile
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree
from typing import Dict, Any

# Cho phép override template từ bên ngoài (dùng khi xuất với template cụ thể)
_OVERRIDE_TEMPLATE = None

def _find_template():
    if _OVERRIDE_TEMPLATE and os.path.exists(_OVERRIDE_TEMPLATE):
        return _OVERRIDE_TEMPLATE
    tpl_dir = "/app/templates"
    for name in ["template.docx", "MẪU_BANG_KÊ_VIỆN_PHÍ_NGOẠI_TRÚ.docx",
                 "MAU_BANG_KE_VIEN_PHI_NGOAI_TRU.docx"]:
        p = os.path.join(tpl_dir, name)
        if os.path.exists(p): return p
    if os.path.isdir(tpl_dir):
        for f in os.listdir(tpl_dir):
            if f.endswith(".docx") and not f.startswith('templates_meta'):
                return os.path.join(tpl_dir, f)
    raise FileNotFoundError(
        "Không tìm thấy file Word mẫu trong /app/templates/. "
        "Hãy upload template qua giao diện Quản lý mẫu in."
    )

# ── Helpers ───────────────────────────────────────────────────────────────────
def fmt(v):
    if v is None: return ''
    try: v = float(v)
    except: return str(v)
    if v == 0: return '0'
    return f"{int(v):,}" if v == int(v) else f"{v:,.2f}"

def fmt2(v):
    if v is None or v == '': return ''
    try:
        f = float(v)
        return f"{f:,.2f}"
    except: return ''

def so_bang_chu(n):
    n = int(round(float(n)))
    if n == 0: return 'Không đồng'
    units = ['','một','hai','ba','bốn','năm','sáu','bảy','tám','chín']
    teens = ['mười','mười một','mười hai','mười ba','mười bốn','mười lăm',
             'mười sáu','mười bảy','mười tám','mười chín']
    def r3(x):
        h,r=divmod(x,100); t,u=divmod(r,10); s=''
        if h: s+=units[h]+' trăm '
        if t==1: s+=teens[u]+' '
        elif t>=2:
            s+=units[t]+' mươi '
            if u==5: s+='lăm '
            elif u==1: s+='mốt '
            elif u: s+=units[u]+' '
        elif u and h: s+='lẻ '+units[u]+' '
        elif u: s+=units[u]+' '
        return s.strip()
    parts=[]
    for name,sc in [('tỷ',1_000_000_000),('triệu',1_000_000),('nghìn',1_000),('',1)]:
        q,n=divmod(n,sc)
        if q: parts.append(r3(q)+(' '+name if name else ''))
    r=' '.join(parts).strip()
    return r[0].upper()+r[1:]+' đồng'

def _add_bold_run(para, text):
    """Thêm run bold vào cuối paragraph."""
    run = para.add_run(text)
    run.bold = True
    return run

def _set_para_value(para, value):
    """
    Tìm vị trí tab cuối cùng trong paragraph rồi ghi value vào run bold sau đó.
    Nếu đã có run bold sau tab thì thay thế, chưa có thì thêm mới.
    """
    runs = para.runs
    last_tab_idx = -1
    for i, r in enumerate(runs):
        if '\t' in r.text:
            last_tab_idx = i
    
    if last_tab_idx < 0:
        # Không có tab, ghi vào cuối
        _add_bold_run(para, value)
        return

    # Xóa các bold run sau tab cuối (nếu có dữ liệu cũ)
    for r in runs[last_tab_idx+1:]:
        if r.bold: r.text = ''
    
    # Thêm run bold mới sau tab
    _add_bold_run(para, value)

def _set_para_after_keyword(para, keyword, value):
    """Ghi value vào run bold ngay sau keyword trong paragraph."""
    runs = para.runs
    found = False
    for i, r in enumerate(runs):
        if keyword in r.text:
            found = True
            # Thêm vào cùng run hoặc run tiếp theo
            if r.bold:
                r.text = r.text.rstrip() + value
            else:
                _add_bold_run(para, value)
            break
    if not found:
        _add_bold_run(para, f' {value}')

def _replace_bold_runs(para, new_text):
    """Thay tất cả bold runs bằng new_text (gộp vào run đầu tiên)."""
    bold_runs = [r for r in para.runs if r.bold and r.text.strip()
                 and r.text.strip() not in ('(đồng)',)]
    if bold_runs:
        bold_runs[0].text = new_text
        for r in bold_runs[1:]: r.text = ''
    else:
        _add_bold_run(para, new_text)

def _set_cell(row, col_idx, text, bold=False, sz=None):
    """Ghi text vào ô bảng. sz: font size in half-points (16=8pt, 18=9pt)."""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    from lxml import etree as _et
    cell = row.cells[col_idx]

    # Xóa toàn bộ runs trong tất cả paragraphs
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''

    para = cell.paragraphs[0]

    def _fix_run(run):
        """Set font size và xóa spacing âm trực tiếp trong XML."""
        rPr = run._r.find(f'{{{W}}}rPr')
        if rPr is None:
            rPr = _et.SubElement(run._r, f'{{{W}}}rPr')
            run._r.insert(0, rPr)
        if bold:
            run.bold = True
        if sz is not None:
            for tag in [f'{{{W}}}sz', f'{{{W}}}szCs']:
                el = rPr.find(tag)
                if el is None:
                    el = _et.SubElement(rPr, tag)
                el.set(f'{{{W}}}val', str(sz))
        # Xóa letter-spacing âm gây wrap sớm
        spc = rPr.find(f'{{{W}}}spacing')
        if spc is not None:
            try:
                if int(spc.get(f'{{{W}}}val', '0')) < 0:
                    rPr.remove(spc)
            except ValueError:
                pass

    if para.runs:
        run = para.runs[0]
        run.text = str(text) if text is not None else ''
        _fix_run(run)
        # Xóa các runs còn lại
        for r in para.runs[1:]:
            r.text = ''
    else:
        run = para.add_run(str(text) if text is not None else '')
        _fix_run(run)

def _set_cell_para(row, col_idx, para_idx, text, bold=False):
    """Ghi text vào paragraph cụ thể trong ô bảng."""
    cell = row.cells[col_idx]
    if para_idx < len(cell.paragraphs):
        para = cell.paragraphs[para_idx]
        bold_runs = [r for r in para.runs if r.bold and r.text.strip()
                     and r.text.strip() not in ('(đồng)',)]
        if bold_runs:
            bold_runs[0].text = text
            for r in bold_runs[1:]: r.text = ''
        else:
            run = para.add_run(text)
            if bold: run.bold = True

# ── Main word generator ───────────────────────────────────────────────────────

def _resize_table_columns(tbl):
    """
    Duyệt trực tiếp XML <w:tc> (không qua python-docx cells vì merged cells gây sai index)
    → set đúng width cho từng cột, xóa spacing âm.
    Font size đã được xử lý trong _set_cell khi ghi data.
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    # Tổng = 11487 dxa. Cột tiền (C3,C4,C6,C8,C9,C12) đủ rộng cho '3,500,000' @ 8pt
    NEW_WIDTHS = [1300, 560, 460, 1050, 1050, 560, 1050, 520, 1000,  980,  980,  980,  997]
    from lxml import etree as _et

    # ── Cập nhật tblGrid (cột master — LibreOffice dùng cái này để render) ──────
    tblGrid = tbl._tbl.find(f'{{{W}}}tblGrid')
    if tblGrid is not None:
        gridCols = tblGrid.findall(f'{{{W}}}gridCol')
        for gc, w in zip(gridCols, NEW_WIDTHS):
            gc.set(f'{{{W}}}w', str(w))


    for tr in tbl._tbl.findall(f'{{{W}}}tr'):
        # Lấy tất cả <w:tc> thực trong row (bỏ qua gridSpan ảo)
        real_tcs = tr.findall(f'{{{W}}}tc')
        # Map từ index tc thực → cột logic (tính gridSpan)
        col_idx = 0
        for tc in real_tcs:
            tcPr = tc.find(f'{{{W}}}tcPr')
            if tcPr is None:
                tcPr = _et.SubElement(tc, f'{{{W}}}tcPr')

            # Tính số cột cell này chiếm (gridSpan)
            gs_el = tcPr.find(f'{{{W}}}gridSpan')
            span = int(gs_el.get(f'{{{W}}}val', '1')) if gs_el is not None else 1

            if span == 1 and col_idx < len(NEW_WIDTHS):
                # Set width
                tcW = tcPr.find(f'{{{W}}}tcW')
                if tcW is None:
                    tcW = _et.SubElement(tcPr, f'{{{W}}}tcW')
                tcW.set(f'{{{W}}}w',    str(NEW_WIDTHS[col_idx]))
                tcW.set(f'{{{W}}}type', 'dxa')
            elif span > 1:
                # Merged cell: set width = sum of spanned columns
                merged_w = sum(NEW_WIDTHS[col_idx:col_idx+span]) if col_idx < len(NEW_WIDTHS) else 0
                if merged_w > 0:
                    tcW = tcPr.find(f'{{{W}}}tcW')
                    if tcW is None:
                        tcW = _et.SubElement(tcPr, f'{{{W}}}tcW')
                    tcW.set(f'{{{W}}}w',    str(merged_w))
                    tcW.set(f'{{{W}}}type', 'dxa')

            col_idx += span

def generate_word(bill: Dict[str, Any], output_path: str):
    TEMPLATE = _find_template()
    doc = Document(TEMPLATE)
    paras = doc.paragraphs

    # ── P1: Mã BN ────────────────────────────────────────────────────────────
    p1 = paras[1]
    # Format: "Mã số người bệnh:\t{ma_bn} Số khám bệnh: {so_kham}"
    for r in p1.runs:
        if r.bold and r.text.strip() and r.text.strip() not in ('Mã', 'số', 'người', 'bệnh:', 'Số', 'khám', 'bệnh:'):
            # Đây là run chứa giá trị mã BN
            if '612939' in r.text:
                r.text = r.text.replace('612939', bill.get('ma_bn', ''))
            break
    # Nếu không tìm thấy run có 612939, thêm run sau tab
    if '612939' not in ' '.join(r.text for r in p1.runs if r.bold):
        # Tìm tab và thêm sau
        for i, r in enumerate(p1.runs):
            if '\t' in r.text and r.bold:
                r.text = '\t' + bill.get('ma_bn', '')
                break
        else:
            _add_bold_run(p1, bill.get('ma_bn', ''))

    # ── P2: Khoa ─────────────────────────────────────────────────────────────
    p2 = paras[2]
    khoa = bill.get('khoa', 'Khoa Khám Bệnh')
    # "Khoa:\t{khoa}" - tab là run bold R3
    tab_found = name_set = False
    for r in p2.runs:
        if not tab_found:
            if '\t' in r.text and r.bold: tab_found = True; r.text = '\t'
        elif not name_set:
            if r.text.strip(): r.text = khoa; name_set = True
            else: r.text = ''
        else: r.text = ''
    if not name_set:
        _add_bold_run(p2, khoa)

    # ── P8: Họ tên / Ngày sinh / Giới tính ───────────────────────────────────
    # Cấu trúc:
    #  R0-R8: "(1) Họ tên người bệnh: "   → chèn họ tên SAU R8, TRƯỚC R9
    #  R9:    TAB bold                     → giữ nguyên
    #  R10-R16: "Ngày, tháng, năm sinh: " → chèn ngày sinh SAU R16, TRƯỚC R18
    #  R18:   TAB bold                     → giữ nguyên
    #  R19-R22: "Giới tính: "             → chèn giới tính SAU R22 (cuối)
    p8 = paras[8]
    ho_ten    = bill.get('ho_ten', '')
    ngay_sinh = bill.get('ngay_sinh', '')
    gioi_tinh = bill.get('gioi_tinh', '')

    from lxml import etree as _et
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _insert_bold_run_after(ref_run, text):
        """Chèn run bold có text ngay sau ref_run."""
        new_r = _et.Element(f'{{{W}}}r')
        rPr = _et.SubElement(new_r, f'{{{W}}}rPr')
        _et.SubElement(rPr, f'{{{W}}}b')
        t = _et.SubElement(new_r, f'{{{W}}}t')
        t.text = text
        # xml:space="preserve" để giữ khoảng trắng
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        ref_run._r.addnext(new_r)

    # Xóa bold run không phải tab (dữ liệu cũ từ lần trước)
    for r in p8.runs:
        if r.bold and '\t' not in r.text:
            r.text = ''

    runs = p8.runs
    # Chèn họ tên sau R8 ("bệnh: "), trước R9 (tab)
    if ho_ten and len(runs) > 8:
        _insert_bold_run_after(runs[8], ho_ten)

    # Sau khi chèn, index runs thay đổi → lấy lại
    runs = p8.runs
    # Tìm R16 ("sinh: ") và chèn ngày sinh sau nó, trước tab thứ 2
    # R16 ban đầu là "sinh:" → sau khi chèn họ tên thì index tăng 1 → R17
    # Tìm run có text "sinh:" rồi chèn sau nó
    for r in runs:
        if r.text.strip() == 'sinh:' and not r.bold:
            _insert_bold_run_after(r, ngay_sinh)
            break

    # Giới tính: thêm vào cuối paragraph
    if gioi_tinh:
        _add_bold_run(p8, gioi_tinh)

    # ── P9: Địa chỉ ──────────────────────────────────────────────────────────
    # Cấu trúc: "(2) Địa chỉ hiện tại: [R8=tại:] [R9=TAB bold] (3) Mã khu vực..."
    # → Địa chỉ phải nằm SAU "tại:" và TRƯỚC tab R9
    p9 = paras[9]
    dia_chi = bill.get('dia_chi', '')
    if dia_chi:
        W_P9 = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        from lxml import etree as _etp9
        for r in p9.runs:
            if r.bold and '\t' not in r.text: r.text = ''
        tab9 = next((r for r in p9.runs if r.bold and '\t' in r.text), None)
        if tab9:
            new_r9 = _etp9.Element(f'{{{W_P9}}}r')
            rPr9 = _etp9.SubElement(new_r9, f'{{{W_P9}}}rPr')
            _etp9.SubElement(rPr9, f'{{{W_P9}}}b')
            t9 = _etp9.SubElement(new_r9, f'{{{W_P9}}}t')
            t9.text = dia_chi
            t9.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            tab9._r.addprevious(new_r9)
        else:
            _add_bold_run(p9, dia_chi)

    # ── P10: Mã BHYT / Giá trị từ / đến ─────────────────────────────────────
    p10 = paras[10]
    ma_bhyt  = bill.get('ma_the_bhyt', '')
    bhyt_tu  = bill.get('bhyt_tu', '')
    bhyt_den = bill.get('bhyt_den', '')
    if ma_bhyt or bhyt_tu or bhyt_den:
        for r in p10.runs:
            if r.bold and '\t' not in r.text: r.text = ''
        if ma_bhyt:  _add_bold_run(p10, ma_bhyt)
        if bhyt_tu:  _add_bold_run(p10, f'\t{bhyt_tu}')
        if bhyt_den: _add_bold_run(p10, f'\t{bhyt_den}')

    # ── P11: Cơ sở đăng ký ───────────────────────────────────────────────────
    p11 = paras[11]
    csdk = bill.get('csdk_bhyt', '')
    if csdk:
        for r in p11.runs:
            if r.bold and '\t' not in r.text: r.text = ''
        _add_bold_run(p11, csdk)

    # ── P12 (body child): (7) Đến khám - ngày khám ───────────────────────────
    # Nằm trong text box (drawing anchor). Có 2 bản sao (AlternateContent).
    # Cấu trúc: "(7) Đến khám: [….] giờ [….] phút, ngày: " → thêm run sau "ngày: "
    ngay_kham = bill.get('ngay_kham', '')
    if ngay_kham:
        W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        body = doc.element.body
        # Tìm body child chứa "(7) Đến khám"
        p_den_kham = None
        for child in list(body):
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                all_text = ''.join(t.text or '' for t in child.findall(f'.//{{{W_NS}}}t'))
                if '(7)' in all_text and 'Đến khám' in all_text:
                    p_den_kham = child
                    break

        if p_den_kham is not None:
            all_runs = p_den_kham.findall(f'.//{{{W_NS}}}r')
            from lxml import etree as _et2
            # "phút, ngày: " là run cuối trước khoảng trống → thêm run bold chứa ngày SAU nó
            # Có 2 bản sao (AlternateContent) nên sẽ có 2 lần cần thêm
            # Nhưng chỉ cần thêm vào bản đầu tiên ((7) ở đầu textbox, R11)
            replaced_count = 0
            for i, r in enumerate(all_runs):
                t_el = r.find(f'{{{W_NS}}}t')
                if t_el is None: continue
                ttext = t_el.text or ''
                if 'phút, ngày:' in ttext and replaced_count == 0:
                    # Thêm run bold mới chứa ngày NGAY SAU run này
                    new_r = _et2.Element(f'{{{W_NS}}}r')
                    rPr_new = _et2.SubElement(new_r, f'{{{W_NS}}}rPr')
                    _et2.SubElement(rPr_new, f'{{{W_NS}}}b')
                    sz = _et2.SubElement(rPr_new, f'{{{W_NS}}}sz')
                    sz.set(f'{{{W_NS}}}val', '20')
                    t_new = _et2.SubElement(new_r, f'{{{W_NS}}}t')
                    t_new.text = ngay_kham
                    t_new.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    r.addnext(new_r)
                    replaced_count += 1
                    break  # chỉ thêm 1 lần (bản đầu tiên)

    # ── P23: Chẩn đoán ───────────────────────────────────────────────────────
    p23 = paras[23]
    chan_doan = bill.get('chan_doan', '')
    ma_benh  = bill.get('ma_benh', '')
    if chan_doan or ma_benh:
        for r in p23.runs:
            if r.bold and '\t' not in r.text: r.text = ''
        if chan_doan: _add_bold_run(p23, chan_doan)
        if ma_benh:  _add_bold_run(p23, f'\t{ma_benh}')

    # ── P26: Bệnh kèm theo ───────────────────────────────────────────────────
    bkt = bill.get('benh_kem_theo', '')
    if bkt:
        p26 = paras[26]
        for r in p26.runs:
            if r.bold and '\t' not in r.text: r.text = ''
        _add_bold_run(p26, bkt)

    # ── P31: Mã BHYT (phần II) / mức hưởng ──────────────────────────────────
    p31 = paras[31]
    muc_huong = bill.get('muc_huong_bhyt', '')
    if ma_bhyt or muc_huong:
        for r in p31.runs:
            if r.bold and '\t' not in r.text: r.text = ''
        if ma_bhyt:  _add_bold_run(p31, ma_bhyt)
        if bhyt_tu:  _add_bold_run(p31, f'\t{bhyt_tu}')
        if bhyt_den: _add_bold_run(p31, f'\t{bhyt_den}')
        if muc_huong:_add_bold_run(p31, f'\t{muc_huong}')

    # ── P32: (Chi phí KBCB tính từ ngày……………….) → điền ngày khám ──────────
    # R11 là "………………." → thay bằng ngày khám
    p32 = paras[32]
    _ngay = bill.get('ngay_kham', '')
    if _ngay:
        for r in p32.runs:
            if '…' in r.text or r.text.count('.') >= 3:
                r.text = ' ' + _ngay
                break
        else:
            for r in p32.runs:
                if r.text.strip() == 'ngày':
                    r.text = f'ngày {_ngay}'
                    break

    # ── Bảng chi phí (Table 0) ────────────────────────────────────────────────
    tbl = doc.tables[0]
    HEADER_ROWS = 3  # R0, R1, R2 = headers

    # Lưu template rows
    tr_group    = copy.deepcopy(tbl.rows[3]._tr)   # "1. Xét nghiệm"
    tr_item     = copy.deepcopy(tbl.rows[4]._tr)   # item row
    tr_subtotal = copy.deepcopy(tbl.rows[14]._tr)  # "Tổng (Xét nghiệm)"
    tr_grand    = copy.deepcopy(tbl.rows[31]._tr)  # "TỔNG CỘNG"
    tr_empty    = copy.deepcopy(tbl.rows[32]._tr)  # empty row
    tr_summary  = copy.deepcopy(tbl.rows[33]._tr)  # summary row

    # Xóa tất cả data rows, giữ header
    while len(tbl.rows) > HEADER_ROWS:
        tbl._tbl.remove(tbl.rows[HEADER_ROWS]._tr)

    def add_group(label):
        tbl._tbl.append(copy.deepcopy(tr_group))
        _set_cell(tbl.rows[-1], 0, label, bold=True)

    def add_item(item):
        tbl._tbl.append(copy.deepcopy(tr_item))
        row = tbl.rows[-1]
        N = 16  # 8pt cho cột số
        _set_cell(row, 0,  item.get('name',''))
        _set_cell(row, 1,  item.get('unit','Lần'),                                             sz=N)
        _set_cell(row, 2,  str(item.get('so_luong',1)),                                        sz=N)
        _set_cell(row, 3,  fmt(item.get('don_gia_bv',0)),                                      sz=N)
        _set_cell(row, 4,  fmt(item.get('don_gia_bh',0)) if item.get('don_gia_bh') else '',    sz=N)
        _set_cell(row, 5,  str(item.get('ty_le_dv',100)),                                      sz=N)
        _set_cell(row, 6,  fmt(item.get('thanh_tien_bv',0)),                                   sz=N)
        _set_cell(row, 7,  str(item.get('ty_le_bh',''))    if item.get('ty_le_bh') else '',    sz=N)
        _set_cell(row, 8,  fmt(item.get('thanh_tien_bh'))  if item.get('thanh_tien_bh') else '',sz=N)
        _set_cell(row, 9,  fmt(item.get('quy_bhyt'))       if item.get('quy_bhyt')  else '',   sz=N)
        _set_cell(row, 10, fmt(item.get('nb_cung_tt'))     if item.get('nb_cung_tt') else '',  sz=N)
        _set_cell(row, 11, fmt(item.get('khac'))           if item.get('khac') else '',        sz=N)
        _set_cell(row, 12, fmt(item.get('nb_tu_tra',0))    if item.get('nb_tu_tra') else '',   sz=N)

    def add_subtotal(gname, bv, bh, quy, nb_tu):
        tbl._tbl.append(copy.deepcopy(tr_subtotal))
        row = tbl.rows[-1]
        N = 16
        _set_cell(row, 0,  f'Tổng ({gname})', bold=True)
        _set_cell(row, 6,  fmt(bv)    if bv    else '', bold=True, sz=N)
        _set_cell(row, 8,  fmt(bh)    if bh    else '', bold=True, sz=N)
        _set_cell(row, 9,  fmt(quy)   if quy   else '', bold=True, sz=N)
        _set_cell(row, 12, fmt(nb_tu) if nb_tu else '', bold=True, sz=N)

    for grp in bill.get('groups', []):
        add_group(f"{grp.get('group_prefix','')}{grp['group_name']}")
        for item in grp.get('items', []): add_item(item)
        add_subtotal(
            grp['group_name'],
            grp.get('tong_tt_bv', 0), grp.get('tong_bh'),
            grp.get('tong_quy'), grp.get('tong_nb_tu_tra'),
        )

    # Grand total
    tbl._tbl.append(copy.deepcopy(tr_grand))
    row = tbl.rows[-1]
    total       = bill.get('tong_cong_bv', 0)
    tong_bh     = bill.get('tong_bh', 0) or 0
    tong_quy    = bill.get('tong_quy', 0) or 0
    tong_nb_tu  = bill.get('tong_nb_tu_tra', total) or total
    _set_cell(row, 0,  'TỔNG CỘNG', bold=True)
    _set_cell(row, 6,  fmt(total),    bold=True)
    _set_cell(row, 8,  fmt(tong_bh)   if tong_bh   else '', bold=True)
    _set_cell(row, 9,  fmt(tong_quy)  if tong_quy  else '', bold=True)
    _set_cell(row, 12, fmt(tong_nb_tu) if tong_nb_tu else '', bold=True)

    # Empty row
    tbl._tbl.append(copy.deepcopy(tr_empty))

    # Summary row (R33) — all content in cell paragraphs
    tbl._tbl.append(copy.deepcopy(tr_summary))
    sum_row  = tbl.rows[-1]
    sum_cell = sum_row.cells[0]
    tong_cp  = bill.get('tong_chi_phi', total)
    quy_bh   = bill.get('quy_bhyt_tt', 0) or 0
    nb_tra   = bill.get('nb_tu_tra', tong_cp) or tong_cp
    khac_tt  = bill.get('khac_tt', 0) or 0
    chu      = so_bang_chu(tong_cp)

    paras_sum = sum_cell.paragraphs
    # P0: Tổng chi phí số
    if len(paras_sum) > 0:
        _replace_bold_runs(paras_sum[0], fmt2(tong_cp))
    # P1: Viết bằng chữ
    if len(paras_sum) > 1:
        after_chu = False; first = False
        for r in paras_sum[1].runs:
            if 'chữ:' in r.text: after_chu = True; continue
            if after_chu:
                if r.text.strip() and r.text.strip() != ')':
                    if not first: r.text = chu + ' '; first = True
                    else: r.text = ''
    # P3: Quỹ BHYT
    if len(paras_sum) > 3:
        _replace_bold_runs(paras_sum[3], fmt2(quy_bh) if quy_bh else '')
    # P4: Người bệnh trả
    if len(paras_sum) > 4:
        _replace_bold_runs(paras_sum[4], fmt2(nb_tra) if nb_tra else '')
    # P5: Các khoản khác
    if len(paras_sum) > 5:
        b = [r for r in paras_sum[5].runs if r.bold and r.text.strip()
             and r.text.strip() not in ('(đồng)',)]
        if b:
            b[0].text = fmt2(khac_tt) if khac_tt else '0.00'
            for r in b[1:]: r.text = ''


    # ── Điều chỉnh độ rộng cột & xóa spacing âm ────────────────────────────
    _resize_table_columns(tbl)

    # ── Fix logo đè tiêu đề (LibreOffice) ───────────────────────────────────
    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    for anchor in doc.element.body.findall('.//{%s}anchor' % WP_NS):
        if anchor.get('behindDoc') == '1':
            anchor.set('relativeHeight', '1')

    doc.save(output_path)
    return output_path


def generate_pdf(bill: Dict[str, Any], output_path: str):
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, 'bill.docx')
        generate_word(bill, docx_path)
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf',
             '--outdir', tmp, docx_path],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice lỗi: {result.stderr}")
        pdf_tmp = os.path.join(tmp, 'bill.pdf')
        if not os.path.exists(pdf_tmp):
            raise RuntimeError("LibreOffice không tạo được PDF")
        shutil.copy2(pdf_tmp, output_path)
    return output_path
